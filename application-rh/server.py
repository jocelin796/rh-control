from __future__ import annotations

import base64
import hashlib
import hmac
import json
import os
import shutil
import sqlite3
import sys
import time
import unicodedata
import uuid
import webbrowser
from datetime import date, datetime, timedelta
from email.message import EmailMessage
from io import BytesIO
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
import smtplib
from typing import Any
from urllib.parse import urlparse

try:
    import psycopg
except ImportError:  # PostgreSQL reste optionnel pour l'utilisation locale SQLite.
    psycopg = None

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.utils.datetime import from_excel
except ImportError:  # L'import Excel reste optionnel jusqu'au build Render.
    Workbook = None
    load_workbook = None
    from_excel = None

try:
    from docx import Document
except ImportError:  # La génération Word reste optionnelle jusqu'au build Render.
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError:  # La génération PDF reste optionnelle jusqu'au build Render.
    A4 = None
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None

ROOT = Path(__file__).resolve().parent
LOCAL_DB_DIR = ROOT / "database"
SCHEMA_PATH = LOCAL_DB_DIR / "schema.sql"
RENDER_DB_PATH = Path("/var/data/rh_control.sqlite")
DATABASE_URL = (
    os.environ.get("DATABASE_URL")
    or os.environ.get("POSTGRES_URL")
    or os.environ.get("SUPABASE_DB_URL")
    or os.environ.get("NEON_DATABASE_URL")
)
DB_MODE = "postgres" if DATABASE_URL else "sqlite"
DB_PATH = Path(os.environ.get("RH_DATABASE_PATH") or (RENDER_DB_PATH if os.environ.get("RENDER") else LOCAL_DB_DIR / "rh_control.sqlite"))
DB_DIR = DB_PATH.parent
HOST = os.environ.get("HOST") or ("0.0.0.0" if os.environ.get("PORT") else "127.0.0.1")
PORT = int(os.environ.get("PORT", "8750"))
APP_PASSWORD = os.environ.get("APP_PASSWORD", "")
APP_SESSION_SECRET = os.environ.get("APP_SESSION_SECRET") or APP_PASSWORD or "local-dev-session-secret"
SESSION_COOKIE = "rh_session"
EMPLOYEE_SESSION_COOKIE = "rh_employee"
SESSION_MAX_AGE_SECONDS = 7 * 24 * 60 * 60
ROLE_PASSWORDS = {
    "Admin RH": APP_PASSWORD,
    "Assistant RH": os.environ.get("ASSISTANT_RH_PASSWORD", ""),
    "Direction": os.environ.get("DIRECTION_PASSWORD", ""),
}
ROLE_PERMISSIONS = {
    "Admin RH": {
        "state_write", "import_json", "export_json", "backup", "excel_import", "excel_template",
        "leave_to_direction", "leave_modify", "leave_refuse", "leave_approve", "document_render",
        "users_manage", "export_excel", "template_import",
    },
    "Assistant RH": {
        "state_write", "excel_import", "excel_template",
        "leave_to_direction", "leave_modify", "leave_refuse", "document_render", "export_excel",
    },
    "Direction": {
        "leave_refuse", "leave_approve", "document_render", "export_excel",
    },
}
SMTP_HOST = os.environ.get("SMTP_HOST", "")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")
SMTP_FROM = os.environ.get("SMTP_FROM") or SMTP_USER
NOTIFY_RH_EMAIL = os.environ.get("NOTIFY_RH_EMAIL", "")
NOTIFY_DIRECTION_EMAIL = os.environ.get("NOTIFY_DIRECTION_EMAIL", "")
EMPLOYEE_EXCEL_COLUMNS = [
    "Noms et prénoms",
    "Matricule",
    "Date d’embauche",
    "Date de naissance",
    "Numéro de téléphone",
    "Adresse mail",
    "Numéro CNPS",
    "Type de contrat",
    "Fonctions",
    "Ville de fonction",
    "Département",
    "Situation matrimoniale",
    "Nombre d’enfants",
    "Date de début de contrat en cours",
    "Date de fin de contrat en cours",
    "Soldes de congé à date",
    "Solde de congé déjà pris",
    "Date de départ",
    "Salaire",
]
DOCUMENT_TYPES = [
    "Attestation de travail",
    "Domiciliation de salaire",
    "Bulletin de salaire",
    "Fiche de congé",
    "Attestation de départ en congé annuel",
    "Certificat de travail",
    "Autres",
]
DEFAULT_DOCUMENT_TEMPLATES = {
    "Attestation de travail": """ATTESTATION DE TRAVAIL

Je soussigné(e), Responsable des Ressources Humaines, atteste que {{nom_complet}}, matricule {{matricule}}, occupe la fonction de {{fonction}} au sein du département {{departement}}.

Type de contrat : {{type_contrat}}
Date d'embauche : {{date_embauche}}
Ville de fonction : {{ville_fonction}}

Fait le {{date_jour}}.
""",
    "Domiciliation de salaire": """DOMICILIATION DE SALAIRE

Nous attestons que {{nom_complet}}, matricule {{matricule}}, est employé(e) en qualité de {{fonction}}.

Salaire de référence : {{salaire}}
Département : {{departement}}

Fait le {{date_jour}}.
""",
    "Bulletin de salaire": """DEMANDE DE BULLETIN DE SALAIRE

Collaborateur : {{nom_complet}}
Matricule : {{matricule}}
Fonction : {{fonction}}
Période / précision : {{precision}}

Fait le {{date_jour}}.
""",
    "Fiche de congé": """FICHE DE CONGÉ

Collaborateur : {{nom_complet}}
Matricule : {{matricule}}
Fonction : {{fonction}}
Département : {{departement}}
Solde de congé à date : {{solde_conge}} jour(s)
Solde déjà pris : {{conge_pris}} jour(s)

Précision : {{precision}}
Fait le {{date_jour}}.
""",
    "Attestation de départ en congé annuel": """ATTESTATION DE DÉPART EN CONGÉ ANNUEL

Nous attestons que {{nom_complet}}, matricule {{matricule}}, part en congé annuel selon les informations validées par l'administration RH.

Fonction : {{fonction}}
Département : {{departement}}
Précision : {{precision}}

Fait le {{date_jour}}.
""",
    "Certificat de travail": """CERTIFICAT DE TRAVAIL

Nous certifions que {{nom_complet}}, matricule {{matricule}}, a travaillé au sein de l'entreprise.

Date d'embauche : {{date_embauche}}
Date de départ : {{date_depart}}
Fonction : {{fonction}}
Département : {{departement}}

Fait le {{date_jour}}.
""",
    "Autres": """DOCUMENT RH

Collaborateur : {{nom_complet}}
Matricule : {{matricule}}
Objet / précision : {{precision}}

Fait le {{date_jour}}.
""",
}
DEFAULT_SETTINGS = {
    "contractAlertDays": [90, 60, 30, 15, 7],
    "returnAlertDays": [3, 1, 0],
    "workingDays": [1, 2, 3, 4, 5],
    "holidays": ["2026-01-01", "2026-04-06", "2026-05-01", "2026-08-07", "2026-12-25"],
    "allowExceptionalLeave": True,
    "ticketEnabled": True,
    "notificationEmails": {
        "rh": "",
        "direction": "",
    },
}


def default_state() -> dict:
    return {
        "currentRole": "Admin RH",
        "settings": dict(DEFAULT_SETTINGS),
        "employees": [],
        "leaveRequests": [],
        "documentRequests": [],
        "documents": [],
        "documentTemplates": dict(DEFAULT_DOCUMENT_TEMPLATES),
        "notifications": [],
        "personnelActions": [],
        "users": [],
        "auditLog": [],
    }


def normalize_state(value: Any) -> dict:
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except json.JSONDecodeError:
            value = {}
    if not isinstance(value, dict):
        value = {}

    state = default_state()
    state["currentRole"] = value.get("currentRole") or state["currentRole"]

    settings = dict(DEFAULT_SETTINGS)
    incoming_settings = value.get("settings")
    if isinstance(incoming_settings, dict):
        settings.update(incoming_settings)
        if isinstance(incoming_settings.get("notificationEmails"), dict):
            emails = dict(DEFAULT_SETTINGS["notificationEmails"])
            emails.update(incoming_settings.get("notificationEmails") or {})
            settings["notificationEmails"] = emails
    state["settings"] = settings

    templates = dict(DEFAULT_DOCUMENT_TEMPLATES)
    incoming_templates = value.get("documentTemplates")
    if isinstance(incoming_templates, dict):
        templates.update({str(k): str(v) for k, v in incoming_templates.items()})
    state["documentTemplates"] = templates

    for key in ["employees", "leaveRequests", "documentRequests", "documents", "notifications", "personnelActions", "users", "auditLog"]:
        items = value.get(key)
        state[key] = items if isinstance(items, list) else []

    return state


def connect() -> sqlite3.Connection:
    if DB_MODE != "sqlite":
        raise RuntimeError("La connexion SQLite est appelée alors que PostgreSQL est configuré.")
    DB_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def connect_postgres():
    if psycopg is None:
        raise RuntimeError("La dépendance PostgreSQL n'est pas installée. Vérifiez requirements.txt et relancez le build.")
    return psycopg.connect(DATABASE_URL)


def init_db() -> None:
    if DB_MODE == "postgres":
        init_postgres_db()
        return

    with connect() as conn:
        schema = SCHEMA_PATH.read_text(encoding="utf-8")
        conn.executescript(schema)
        conn.execute(
            "INSERT OR REPLACE INTO app_meta(key, value) VALUES (?, ?)",
            ("schema_version", "1"),
        )


def init_postgres_db() -> None:
    with connect_postgres() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS app_state (
                  key TEXT PRIMARY KEY,
                  value JSONB NOT NULL,
                  updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """
            )
            cur.execute(
                """
                INSERT INTO app_state(key, value)
                VALUES (%s, %s::jsonb)
                ON CONFLICT (key) DO NOTHING
                """,
                ("main", json.dumps(default_state(), ensure_ascii=False)),
            )


def db_is_durable() -> bool:
    if DB_MODE == "postgres":
        return True
    path = str(DB_PATH).replace("\\", "/")
    if os.environ.get("RENDER"):
        return path.startswith("/var/data/")
    return True


def database_label() -> str:
    if DB_MODE == "postgres":
        return "PostgreSQL externe (DATABASE_URL)"
    return str(DB_PATH)


def backup_root() -> Path:
    if DB_MODE == "postgres" and os.environ.get("RENDER"):
        return Path("/tmp/rh-control-backups")
    return DB_DIR / "backups"


def normalize_lookup(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = "".join(
        char for char in unicodedata.normalize("NFKD", text)
        if not unicodedata.combining(char)
    )
    return "".join(char for char in text if char.isalnum())


def find_employee_by_matricule(data: dict, matricule: str) -> dict | None:
    target = normalize_lookup(matricule)
    for employee in data.get("employees", []):
        if normalize_lookup(employee.get("matricule")) == target:
            return employee
    return None


def find_employee_by_id(data: dict, employee_id: str) -> dict | None:
    for employee in data.get("employees", []):
        if employee.get("id") == employee_id:
            return employee
    return None


def parse_excel_date(value: Any) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, (int, float)) and from_excel is not None:
        try:
            return from_excel(value).date().isoformat()
        except Exception:
            return ""
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except ValueError:
            continue
    return text


def parse_number(value: Any) -> float:
    if value is None or value == "":
        return 0
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip().replace("\u00a0", "").replace(" ", "").replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return 0


def split_full_name(full_name: str) -> tuple[str, str]:
    parts = str(full_name or "").strip().split()
    if not parts:
        return "", ""
    if len(parts) == 1:
        return parts[0], ""
    last_name = parts[0]
    first_name = " ".join(parts[1:])
    return first_name, last_name


def row_get(row: dict[str, Any], *aliases: str) -> Any:
    for alias in aliases:
        key = normalize_lookup(alias)
        if key in row and row[key] not in (None, ""):
            return row[key]
    return ""


def employee_from_excel_row(row: dict[str, Any], existing: dict | None = None) -> dict | None:
    matricule = str(row_get(row, "Matricule", "Matricule Zeus")).strip()
    if not matricule:
        return None

    full_name = str(row_get(row, "Noms et prénoms", "Nom et prénoms", "Nom complet", "Collaborateur")).strip()
    first_name = str(row_get(row, "Prénom", "Prénoms", "Prenom", "Prenoms")).strip()
    last_name = str(row_get(row, "Nom", "Noms")).strip()
    if full_name and (not first_name or not last_name):
        first_name, last_name = split_full_name(full_name)

    hire_date = parse_excel_date(row_get(row, "Date d’embauche", "Date d'embauche", "Date embauche"))
    birth_date = parse_excel_date(row_get(row, "Date de naissance", "Naissance"))
    contract_start = parse_excel_date(row_get(row, "Date de début de contrat en cours", "Date debut contrat", "Début contrat"))
    contract_end = parse_excel_date(row_get(row, "Date de fin de contrat en cours", "Date fin contrat", "Fin contrat"))
    departure_date = parse_excel_date(row_get(row, "Date de départ", "Date depart", "Départ"))

    leave_available = parse_number(row_get(row, "Soldes de congé à date", "Solde de congé à date", "Solde congé à date", "Solde conge"))
    leave_taken = parse_number(row_get(row, "Solde de congé déjà pris", "Congé déjà pris", "Conge deja pris", "Congés pris"))
    initial_leave = leave_available + leave_taken if leave_taken else leave_available

    service = str(row_get(row, "Département", "Departement", "Service")).strip()
    city = str(row_get(row, "Ville de fonction", "Ville", "Agence")).strip()
    fonction = str(row_get(row, "Fonctions", "Fonction", "Poste")).strip()
    contract_type = str(row_get(row, "Type de contrat", "Contrat")).strip() or "Autre"
    salary = parse_number(row_get(row, "Salaire", "Rémunération", "Remuneration"))

    employee_id = existing.get("id") if existing else f"emp-{uuid.uuid4().hex[:12]}"
    contract_id = ""
    existing_contracts = existing.get("contracts", []) if existing else []
    if existing_contracts:
        contract_id = existing_contracts[-1].get("id", "")
    if not contract_id:
        contract_id = f"ctr-{uuid.uuid4().hex[:12]}"

    previous_history = []
    if existing_contracts:
        previous_history = existing_contracts[-1].get("history", [])

    employee = {
        "id": employee_id,
        "matricule": matricule,
        "firstName": first_name,
        "lastName": last_name,
        "hireDate": hire_date,
        "birthDate": birth_date,
        "phone": str(row_get(row, "Numéro de téléphone", "Numero de telephone", "Téléphone", "Telephone")).strip(),
        "email": str(row_get(row, "Adresse mail", "Email", "Adresse email", "Mail")).strip(),
        "cnpsNumber": str(row_get(row, "Numéro CNPS", "Numero CNPS", "CNPS")).strip(),
        "service": service,
        "direction": existing.get("direction", "") if existing else service,
        "agency": city,
        "fonction": fonction,
        "maritalStatus": str(row_get(row, "Situation matrimoniale", "Situations matrimoniales", "Situation familiale")).strip(),
        "childrenCount": int(parse_number(row_get(row, "Nombre d’enfants", "Nombre enfants", "Enfants"))),
        "departureDate": departure_date,
        "status": "Sorti" if departure_date else "Actif",
        "leaveBalance": {
            "initial": initial_leave,
            "acquired": 0,
            "taken": leave_taken,
            "planned": existing.get("leaveBalance", {}).get("planned", 0) if existing else 0,
            "available": leave_available,
        },
        "contracts": [
            {
                "id": contract_id,
                "type": contract_type,
                "start": contract_start or hire_date,
                "end": contract_end,
                "duration": "",
                "fonction": fonction,
                "service": service,
                "salary": salary,
                "renewalDate": existing_contracts[-1].get("renewalDate", "") if existing_contracts else "",
                "renewalCount": existing_contracts[-1].get("renewalCount", 0) if existing_contracts else 0,
                "document": existing_contracts[-1].get("document", "") if existing_contracts else "",
                "status": "Contrat actif",
                "history": previous_history + [f"{datetime.now():%d/%m/%Y} : fiche mise à jour par import Excel."],
            }
        ],
    }
    return employee


def workbook_rows_from_excel(content: bytes) -> tuple[list[str], list[tuple[int, dict[str, Any]]]]:
    if load_workbook is None:
        raise RuntimeError("La lecture Excel n'est pas disponible. Vérifiez la dépendance openpyxl.")

    workbook = load_workbook(filename=BytesIO(content), data_only=True)
    sheet = workbook.active
    header_cells = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_cells:
        raise RuntimeError("Le fichier Excel ne contient pas d'en-têtes.")

    headers = [normalize_lookup(value) for value in header_cells]
    parsed_rows: list[tuple[int, dict[str, Any]]] = []
    for row_number, excel_row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        if not any(cell not in (None, "") for cell in excel_row):
            continue
        row = {headers[index]: excel_row[index] for index in range(min(len(headers), len(excel_row)))}
        parsed_rows.append((row_number, row))
    return headers, parsed_rows


def preview_employees_from_excel(content: bytes) -> dict:
    headers, parsed_rows = workbook_rows_from_excel(content)
    data = get_state()
    existing_by_matricule = {
        normalize_lookup(employee.get("matricule")): employee
        for employee in data.get("employees", [])
    }
    missing_columns = [
        column for column in EMPLOYEE_EXCEL_COLUMNS
        if normalize_lookup(column) not in headers
    ]
    results = []
    created = 0
    updated = 0
    skipped = 0
    errors_count = 0

    for row_number, row in parsed_rows:
        matricule = str(row_get(row, "Matricule", "Matricule Zeus")).strip()
        full_name = str(row_get(row, "Noms et prénoms", "Nom et prénoms", "Nom complet", "Collaborateur")).strip()
        first_name = str(row_get(row, "Prénom", "Prénoms", "Prenom", "Prenoms")).strip()
        last_name = str(row_get(row, "Nom", "Noms")).strip()
        row_errors = []
        row_warnings = []
        if not matricule:
            row_errors.append("Matricule Zeus manquant")
            skipped += 1
        if not (full_name or first_name or last_name):
            row_errors.append("Nom et prénoms manquants")
        contract_end = parse_excel_date(row_get(row, "Date de fin de contrat en cours", "Date fin contrat", "Fin contrat"))
        contract_start = parse_excel_date(row_get(row, "Date de début de contrat en cours", "Date debut contrat", "Début contrat"))
        if contract_start and contract_end and contract_end < contract_start:
            row_warnings.append("La date de fin de contrat est avant la date de début")
        key = normalize_lookup(matricule)
        existing = existing_by_matricule.get(key)
        action = "Erreur" if row_errors else ("Mise à jour" if existing else "Création")
        if row_errors:
            errors_count += 1
            skipped += 1 if matricule else 0
        elif existing:
            updated += 1
        else:
            created += 1
        results.append({
            "row": row_number,
            "matricule": matricule,
            "name": full_name or f"{first_name} {last_name}".strip(),
            "department": str(row_get(row, "Département", "Departement", "Service")).strip(),
            "function": str(row_get(row, "Fonctions", "Fonction", "Poste")).strip(),
            "action": action,
            "errors": row_errors,
            "warnings": row_warnings,
        })

    return {
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "errors": errors_count,
        "rows": results,
        "missingColumns": missing_columns,
    }


def import_employees_from_excel(content: bytes, actor: str = "Admin RH") -> dict:
    preview = preview_employees_from_excel(content)
    if preview.get("errors"):
        raise RuntimeError("Le fichier contient des lignes en erreur. Corrige-les ou retire-les avant l'import.")

    _, parsed_rows = workbook_rows_from_excel(content)
    data = get_state()
    existing_by_matricule = {
        normalize_lookup(employee.get("matricule")): employee
        for employee in data.get("employees", [])
    }

    created = 0
    updated = 0
    skipped = 0

    for _, row in parsed_rows:
        matricule = str(row_get(row, "Matricule", "Matricule Zeus")).strip()
        if not matricule:
            skipped += 1
            continue
        full_name = str(row_get(row, "Noms et prénoms", "Nom et prénoms", "Nom complet", "Collaborateur")).strip()
        first_name = str(row_get(row, "Prénom", "Prénoms", "Prenom", "Prenoms")).strip()
        last_name = str(row_get(row, "Nom", "Noms")).strip()
        if not (full_name or first_name or last_name):
            skipped += 1
            continue
        key = normalize_lookup(matricule)
        existing = existing_by_matricule.get(key)
        employee = employee_from_excel_row(row, existing)
        if not employee:
            skipped += 1
            continue
        if existing:
            index = data["employees"].index(existing)
            data["employees"][index] = employee
            updated += 1
        else:
            data["employees"].append(employee)
            existing_by_matricule[key] = employee
            created += 1

    data.setdefault("auditLog", []).insert(0, {
        "date": datetime.now().date().isoformat(),
        "actor": actor,
        "action": f"Import Excel collaborateurs : {created} créé(s), {updated} mis à jour, {skipped} ignoré(s).",
    })
    save_state(data)
    return {"created": created, "updated": updated, "skipped": skipped, "preview": preview, "data": get_state()}


def build_employee_template_workbook() -> bytes:
    if Workbook is None:
        raise RuntimeError("La création du modèle Excel n'est pas disponible.")
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Collaborateurs"
    sheet.append(EMPLOYEE_EXCEL_COLUMNS)
    for index, column in enumerate(EMPLOYEE_EXCEL_COLUMNS, start=1):
        sheet.cell(row=1, column=index).style = "Headline 3"
        sheet.column_dimensions[sheet.cell(row=1, column=index).column_letter].width = max(18, len(column) + 2)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def append_sheet(workbook: Workbook, title: str, headers: list[str], rows_data: list[list[Any]]) -> None:
    sheet = workbook.create_sheet(title=title[:31])
    sheet.append(headers)
    for row in rows_data:
        sheet.append(row)
    for index, column in enumerate(headers, start=1):
        sheet.cell(row=1, column=index).style = "Headline 3"
        sheet.column_dimensions[sheet.cell(row=1, column=index).column_letter].width = max(16, min(34, len(column) + 5))


def build_export_workbook(export_type: str = "summary") -> bytes:
    if Workbook is None:
        raise RuntimeError("La création Excel n'est pas disponible.")
    data = get_state()
    workbook = Workbook()
    workbook.remove(workbook.active)
    export_type = str(export_type or "summary").lower()

    employees = data.get("employees", [])
    leaves = data.get("leaveRequests", [])
    documents = data.get("documentRequests", [])
    personnel_actions = data.get("personnelActions", [])

    if export_type in {"summary", "all"}:
        active = len([employee for employee in employees if employee.get("status") == "Actif"])
        pending_leaves = len([leave for leave in leaves if leave.get("status") in {"Demande envoyée", "Modification demandée", "En attente de validation Direction"}])
        pending_docs = len([request for request in documents if request.get("status") not in {"Document transmis", "Refusé"}])
        expiring = 0
        expired = 0
        for employee in employees:
            contract = (employee.get("contracts") or [{}])[-1]
            end = contract.get("end")
            if not end:
                continue
            try:
                days = (datetime.strptime(end, "%Y-%m-%d").date() - datetime.now().date()).days
            except ValueError:
                continue
            if days < 0:
                expired += 1
            elif days <= 90:
                expiring += 1
        append_sheet(workbook, "Synthèse", ["Indicateur", "Valeur"], [
            ["Collaborateurs", len(employees)],
            ["Collaborateurs actifs", active],
            ["Contrats à échéance ≤ 90 jours", expiring],
            ["Contrats expirés", expired],
            ["Demandes de congé en attente", pending_leaves],
            ["Demandes de document en attente", pending_docs],
            ["Actions personnel", len(personnel_actions)],
        ])

    if export_type in {"employees", "all"}:
        append_sheet(workbook, "Collaborateurs", [
            "Matricule", "Nom", "Prénom", "Fonction", "Département", "Ville", "Téléphone", "Email",
            "Type contrat", "Début contrat", "Fin contrat", "Salaire", "Solde congé disponible", "Statut",
        ], [
            [
                employee.get("matricule", ""),
                employee.get("lastName", ""),
                employee.get("firstName", ""),
                employee.get("fonction", ""),
                employee.get("service", ""),
                employee.get("agency", ""),
                employee.get("phone", ""),
                employee.get("email", ""),
                ((employee.get("contracts") or [{}])[-1]).get("type", ""),
                ((employee.get("contracts") or [{}])[-1]).get("start", ""),
                ((employee.get("contracts") or [{}])[-1]).get("end", ""),
                ((employee.get("contracts") or [{}])[-1]).get("salary", 0),
                (employee.get("leaveBalance") or {}).get("available", 0),
                employee.get("status", ""),
            ]
            for employee in employees
        ])

    if export_type in {"contracts", "all"}:
        rows_data = []
        for employee in employees:
            for contract in employee.get("contracts", []):
                rows_data.append([
                    employee.get("matricule", ""),
                    full_name_server(employee),
                    contract.get("type", ""),
                    contract.get("start", ""),
                    contract.get("end", ""),
                    contract.get("fonction", ""),
                    contract.get("service", ""),
                    contract.get("salary", 0),
                    contract.get("status", ""),
                    contract.get("renewalCount", 0),
                ])
        append_sheet(workbook, "Contrats", ["Matricule", "Collaborateur", "Type", "Début", "Fin", "Fonction", "Service", "Salaire", "Statut", "Renouvellements"], rows_data)

    if export_type in {"leaves", "all"}:
        append_sheet(workbook, "Congés", ["Matricule", "Collaborateur", "Type", "Début", "Fin", "Reprise", "Jours", "Statut", "Commentaire"], [
            [
                (find_employee_by_id(data, leave.get("employeeId", "")) or {}).get("matricule", ""),
                full_name_server(find_employee_by_id(data, leave.get("employeeId", "")) or {}),
                leave.get("type", ""),
                leave.get("start", ""),
                leave.get("end", ""),
                leave.get("returnDate", ""),
                leave.get("days", 0),
                leave.get("status", ""),
                leave.get("comment", ""),
            ]
            for leave in leaves
        ])

    if export_type in {"documents", "all"}:
        append_sheet(workbook, "Documents", ["Matricule", "Collaborateur", "Document", "Date demande", "Statut", "Détail"], [
            [
                (find_employee_by_id(data, request.get("employeeId", "")) or {}).get("matricule", ""),
                full_name_server(find_employee_by_id(data, request.get("employeeId", "")) or {}),
                request.get("type", ""),
                request.get("createdAt", ""),
                request.get("status", ""),
                request.get("details", ""),
            ]
            for request in documents
        ])

    if export_type in {"personnel", "all"}:
        append_sheet(workbook, "Personnel", ["Matricule", "Collaborateur", "Type", "Date", "Objet", "Statut", "Détail"], [
            [
                (find_employee_by_id(data, action.get("employeeId", "")) or {}).get("matricule", ""),
                full_name_server(find_employee_by_id(data, action.get("employeeId", "")) or {}),
                action.get("type", ""),
                action.get("date", ""),
                action.get("title", ""),
                action.get("status", ""),
                action.get("details", ""),
            ]
            for action in personnel_actions
        ])

    if not workbook.sheetnames:
        append_sheet(workbook, "Synthèse", ["Indicateur", "Valeur"], [["Aucune donnée", 0]])

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def import_docx_template(content: bytes) -> str:
    if Document is None:
        raise RuntimeError("La lecture Word n'est pas disponible. Vérifiez la dépendance python-docx.")
    document = Document(BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    text = "\n".join(lines).strip()
    if not text:
        raise RuntimeError("Le modèle Word ne contient pas de texte lisible.")
    return text


def extract_multipart_file(body: bytes, content_type: str) -> bytes:
    marker = "boundary="
    if marker not in content_type:
        raise RuntimeError("Fichier Excel introuvable.")
    boundary = content_type.split(marker, 1)[1].strip().strip('"')
    boundary_bytes = ("--" + boundary).encode("utf-8")
    for part in body.split(boundary_bytes):
        if b"Content-Disposition" not in part or b'name=\"file\"' not in part:
            continue
        if b"\r\n\r\n" not in part:
            continue
        _, payload = part.split(b"\r\n\r\n", 1)
        payload = payload.strip(b"\r\n")
        if payload.endswith(b"--"):
            payload = payload[:-2].strip(b"\r\n")
        return payload
    raise RuntimeError("Fichier Excel introuvable.")


def calc_leave_days_server(settings: dict, start: str, end: str) -> int:
    try:
        start_date = datetime.strptime(start, "%Y-%m-%d").date()
        end_date = datetime.strptime(end, "%Y-%m-%d").date()
    except ValueError:
        return 0
    if end_date < start_date:
        return 0
    working_days = set(settings.get("workingDays", [1, 2, 3, 4, 5]))
    holidays = set(settings.get("holidays", []))
    count = 0
    current = start_date
    while current <= end_date:
        js_weekday = (current.weekday() + 1) % 7
        if js_weekday in working_days and current.isoformat() not in holidays:
            count += 1
        current += timedelta(days=1)
    return count


def calc_return_date_server(settings: dict, end: str) -> str:
    try:
        current = datetime.strptime(end, "%Y-%m-%d").date() + timedelta(days=1)
    except ValueError:
        return ""
    working_days = set(settings.get("workingDays", [1, 2, 3, 4, 5]))
    holidays = set(settings.get("holidays", []))
    while True:
        js_weekday = (current.weekday() + 1) % 7
        if js_weekday in working_days and current.isoformat() not in holidays:
            return current.isoformat()
        current += timedelta(days=1)


def employee_portal_payload(data: dict, employee: dict) -> dict:
    employee_id = employee.get("id")
    return {
        "employee": employee,
        "settings": data.get("settings", {}),
        "leaveRequests": [
            leave for leave in data.get("leaveRequests", [])
            if leave.get("employeeId") == employee_id
        ],
        "documentRequests": [
            request for request in data.get("documentRequests", [])
            if request.get("employeeId") == employee_id
        ],
        "documents": [
            document for document in data.get("documents", [])
            if document.get("employeeId") == employee_id
        ],
    }


def full_name_server(employee: dict) -> str:
    return " ".join(
        part for part in [employee.get("firstName", ""), employee.get("lastName", "")]
        if str(part).strip()
    ).strip() or employee.get("matricule", "Collaborateur")


def format_date_fr(value: str) -> str:
    if not value:
        return "—"
    try:
        return datetime.strptime(value, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return value


def today_iso() -> str:
    return datetime.now().date().isoformat()


def make_leave_ticket_server(leave: dict, employee: dict) -> str:
    observations = leave.get("observations") or []
    return f"""TICKET DE CONGÉ

Collaborateur : {full_name_server(employee)}
Matricule : {employee.get('matricule', '')}
Fonction : {employee.get('fonction', '')}
Service : {employee.get('service', '')}

Type de congé : {leave.get('type', '')}
Date de départ : {format_date_fr(leave.get('start', ''))}
Date de fin : {format_date_fr(leave.get('end', ''))}
Date de reprise : {format_date_fr(leave.get('returnDate', ''))}
Nombre de jours : {leave.get('days', 0)}

Statut : {leave.get('status', '')}
Généré le : {format_date_fr(today_iso())}

Observation RH / Direction :
{chr(10).join(observations) if observations else 'Aucune observation.'}"""


def record_notification(data: dict, event_type: str, subject: str, body: str, recipients: list[str] | None = None) -> dict:
    recipients = [email for email in (recipients or []) if email]
    item = {
        "id": f"notif-{uuid.uuid4().hex[:12]}",
        "date": datetime.now().isoformat(timespec="seconds"),
        "type": event_type,
        "subject": subject,
        "body": body,
        "recipients": recipients,
        "status": "Journalisé",
        "error": "",
    }
    if SMTP_HOST and SMTP_FROM and recipients:
        try:
            message = EmailMessage()
            message["Subject"] = subject
            message["From"] = SMTP_FROM
            message["To"] = ", ".join(recipients)
            message.set_content(body)
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=12) as smtp:
                smtp.starttls()
                if SMTP_USER and SMTP_PASSWORD:
                    smtp.login(SMTP_USER, SMTP_PASSWORD)
                smtp.send_message(message)
            item["status"] = "E-mail envoyé"
            item["sentAt"] = datetime.now().isoformat(timespec="seconds")
        except Exception as exc:
            item["status"] = "Échec e-mail"
            item["error"] = str(exc)
    data.setdefault("notifications", []).insert(0, item)
    return item


def notification_recipients(data: dict, target: str) -> list[str]:
    emails = ((data.get("settings") or {}).get("notificationEmails") or {})
    if target == "rh":
        return [emails.get("rh") or NOTIFY_RH_EMAIL]
    if target == "direction":
        return [emails.get("direction") or NOTIFY_DIRECTION_EMAIL]
    return []


def notify_employee_if_possible(data: dict, employee: dict, subject: str, body: str, event_type: str) -> None:
    record_notification(data, event_type, subject, body, [employee.get("email", "")])


def transition_leave_server(data: dict, leave_id: str, action: str, role: str, observation: str = "") -> dict:
    permission = f"leave_{action.replace('-', '_')}"
    if not role_has_permission(role, permission):
        raise PermissionError("Action non autorisée pour ce rôle.")

    leave = next((item for item in data.get("leaveRequests", []) if item.get("id") == leave_id), None)
    if not leave:
        raise RuntimeError("Demande de congé introuvable.")
    employee = find_employee_by_id(data, leave.get("employeeId", ""))
    if not employee:
        raise RuntimeError("Collaborateur introuvable.")

    leave.setdefault("observations", [])
    leave.setdefault("history", [])
    actor = normalize_role(role)
    stamp = format_date_fr(today_iso())
    name = full_name_server(employee)
    observation = observation.strip()

    if action == "to-direction":
        leave["status"] = "En attente de validation Direction"
        leave["observations"].append(f"{actor} : {observation or 'demande vérifiée et transmise à la Direction.'}")
        leave["history"].append(f"{stamp} : demande transmise à la Direction.")
        data.setdefault("auditLog", []).insert(0, {"date": today_iso(), "actor": actor, "action": f"Demande de {name} transmise à la Direction."})
        record_notification(
            data,
            "Congés",
            "Demande de congé à valider",
            f"{name} attend une validation Direction pour {leave.get('days', 0)} jour(s), du {format_date_fr(leave.get('start', ''))} au {format_date_fr(leave.get('end', ''))}.",
            notification_recipients(data, "direction"),
        )
    elif action == "modify":
        leave["status"] = "Modification demandée"
        leave["observations"].append(f"{actor} : {observation or 'modification demandée.'}")
        leave["history"].append(f"{stamp} : modification demandée.")
        data.setdefault("auditLog", []).insert(0, {"date": today_iso(), "actor": actor, "action": f"Modification demandée pour le congé de {name}."})
        notify_employee_if_possible(data, employee, "Modification demandée sur votre congé", f"Observation : {observation or 'modification demandée.'}", "Congés")
    elif action == "refuse":
        leave["status"] = "Refusé"
        leave["observations"].append(f"{actor} : {observation or 'demande refusée.'}")
        leave["history"].append(f"{stamp} : demande refusée.")
        data.setdefault("auditLog", []).insert(0, {"date": today_iso(), "actor": actor, "action": f"Demande de congé refusée pour {name}."})
        notify_employee_if_possible(data, employee, "Demande de congé refusée", f"Observation : {observation or 'demande refusée.'}", "Congés")
    elif action == "approve":
        was_validated = leave.get("status") == "Validé"
        leave["status"] = "Validé"
        leave["observations"].append(f"{actor} : {observation or 'demande validée.'}")
        leave["history"].append(f"{stamp} : demande validée par la Direction.")
        if not was_validated:
            balance = employee.setdefault("leaveBalance", {})
            balance["taken"] = float(balance.get("taken", 0) or 0) + float(leave.get("days", 0) or 0)
            balance["available"] = float(balance.get("available", 0) or 0) - float(leave.get("days", 0) or 0)
        if data.get("settings", {}).get("ticketEnabled", True) and not any(doc.get("leaveId") == leave_id for doc in data.get("documents", [])):
            content = make_leave_ticket_server(leave, employee)
            data.setdefault("documents", []).append({
                "id": f"doc-{uuid.uuid4().hex[:12]}",
                "employeeId": employee.get("id", ""),
                "leaveId": leave.get("id", ""),
                "title": f"Ticket de congé - {name}",
                "status": "Document à transmettre",
                "createdAt": today_iso(),
                "content": content,
            })
            leave["history"].append(f"{stamp} : ticket de congé généré.")
        data.setdefault("auditLog", []).insert(0, {"date": today_iso(), "actor": actor, "action": f"Demande de congé validée pour {name}."})
        notify_employee_if_possible(data, employee, "Demande de congé validée", f"Votre congé du {format_date_fr(leave.get('start', ''))} au {format_date_fr(leave.get('end', ''))} est validé.", "Congés")
    else:
        raise RuntimeError("Action congé inconnue.")
    return data


def safe_filename(value: str, suffix: str) -> str:
    text = normalize_lookup(value) or "document_rh"
    return f"{text}.{suffix}"


def xml_escape(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def render_document_file(title: str, content: str, file_format: str) -> tuple[bytes, str, str]:
    title = str(title or "Document RH").strip() or "Document RH"
    content = str(content or "").strip()
    file_format = str(file_format or "docx").lower()
    if file_format == "docx":
        if Document is None:
            raise RuntimeError("La génération Word n'est pas disponible. Vérifiez la dépendance python-docx.")
        document = Document()
        section = document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "PALLADIUM AFRIQUE · RH CONTROL"
        header_para.alignment = 1
        document.add_heading(title, level=1)
        document.add_paragraph(f"Date d'édition : {format_date_fr(today_iso())}")
        document.add_paragraph("")
        for block in content.split("\n\n"):
            lines = [line for line in block.splitlines()]
            if not lines:
                document.add_paragraph("")
            else:
                paragraph = document.add_paragraph()
                for index, line in enumerate(lines):
                    if index:
                        paragraph.add_run().add_break()
                    paragraph.add_run(line)
        document.add_paragraph("")
        document.add_paragraph("Fait pour servir et valoir ce que de droit.")
        document.add_paragraph("Signature et cachet :")
        buffer = BytesIO()
        document.save(buffer)
        return (
            buffer.getvalue(),
            safe_filename(title, "docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if file_format == "pdf":
        if SimpleDocTemplate is None or getSampleStyleSheet is None or Paragraph is None or Spacer is None:
            raise RuntimeError("La génération PDF n'est pas disponible. Vérifiez la dépendance reportlab.")
        buffer = BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=48, leftMargin=48, topMargin=54, bottomMargin=48)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("PALLADIUM AFRIQUE · RH CONTROL", styles["Heading3"]),
            Spacer(1, 8),
            Paragraph(xml_escape(title), styles["Title"]),
            Paragraph(f"Date d'édition : {format_date_fr(today_iso())}", styles["BodyText"]),
            Spacer(1, 18),
        ]
        for block in content.split("\n\n"):
            html = "<br/>".join(xml_escape(line) for line in block.splitlines()) or "&nbsp;"
            story.append(Paragraph(html, styles["BodyText"]))
            story.append(Spacer(1, 9))
        story.extend([
            Spacer(1, 18),
            Paragraph("Fait pour servir et valoir ce que de droit.", styles["BodyText"]),
            Paragraph("Signature et cachet :", styles["BodyText"]),
        ])
        pdf.build(story)
        return buffer.getvalue(), safe_filename(title, "pdf"), "application/pdf"
    raise RuntimeError("Format non reconnu. Utilise docx ou pdf.")


def sanitize_state_for_role(incoming: dict, role: str) -> dict:
    state = normalize_state(incoming)
    current = get_state()
    current_users_by_id = {user.get("id"): user for user in current.get("users", [])}
    for user in state.get("users", []):
        current_user = current_users_by_id.get(user.get("id"))
        if current_user and current_user.get("passwordHash") and not user.get("passwordHash"):
            user["passwordHash"] = current_user.get("passwordHash")
    if normalize_role(role) == "Admin RH":
        return state
    state["settings"] = current.get("settings", dict(DEFAULT_SETTINGS))
    state["documentTemplates"] = current.get("documentTemplates", dict(DEFAULT_DOCUMENT_TEMPLATES))
    state["users"] = current.get("users", [])
    state["currentRole"] = current.get("currentRole", "Admin RH")
    return state


def hash_password(password: str) -> str:
    iterations = 260000
    salt = os.urandom(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return "pbkdf2_sha256${}${}${}".format(
        iterations,
        base64.urlsafe_b64encode(salt).decode("ascii"),
        base64.urlsafe_b64encode(digest).decode("ascii"),
    )


def verify_password_hash(password: str, password_hash: str) -> bool:
    try:
        algorithm, iterations, salt_b64, digest_b64 = str(password_hash or "").split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        salt = base64.urlsafe_b64decode(salt_b64.encode("ascii"))
        expected = base64.urlsafe_b64decode(digest_b64.encode("ascii"))
        actual = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, int(iterations))
        return hmac.compare_digest(actual, expected)
    except Exception:
        return False


def public_user(user: dict) -> dict:
    return {
        "id": user.get("id", ""),
        "username": user.get("username", ""),
        "fullName": user.get("fullName", ""),
        "role": normalize_role(user.get("role")),
        "active": user.get("active", True) is not False,
        "createdAt": user.get("createdAt", ""),
        "updatedAt": user.get("updatedAt", ""),
        "lastLoginAt": user.get("lastLoginAt", ""),
        "hasPassword": bool(user.get("passwordHash")),
    }


def state_for_client(data: dict) -> dict:
    public_state = normalize_state(data)
    public_state["users"] = [public_user(user) for user in data.get("users", [])]
    return public_state


def upsert_user(data: dict, payload: dict, actor: str) -> dict:
    users = data.setdefault("users", [])
    user_id = str(payload.get("id", "")).strip()
    username = str(payload.get("username", "")).strip()
    if not username:
        raise RuntimeError("Identifiant utilisateur obligatoire.")
    role = normalize_role(payload.get("role"))
    password = str(payload.get("password", ""))
    now = datetime.now().isoformat(timespec="seconds")
    existing = next((user for user in users if user.get("id") == user_id), None) if user_id else None
    if not existing:
        for user in users:
            if normalize_lookup(user.get("username")) == normalize_lookup(username):
                raise RuntimeError("Cet identifiant existe déjà.")
        existing = {
            "id": f"user-{uuid.uuid4().hex[:12]}",
            "createdAt": now,
            "passwordHash": "",
        }
        users.append(existing)
    elif normalize_lookup(existing.get("username")) != normalize_lookup(username):
        for user in users:
            if user.get("id") != existing.get("id") and normalize_lookup(user.get("username")) == normalize_lookup(username):
                raise RuntimeError("Cet identifiant existe déjà.")

    if not existing.get("passwordHash") and not password:
        raise RuntimeError("Mot de passe obligatoire pour un nouvel utilisateur.")

    existing.update({
        "username": username,
        "fullName": str(payload.get("fullName", "")).strip() or username,
        "role": role,
        "active": payload.get("active", True) is not False,
        "updatedAt": now,
    })
    if password:
        if len(password) < 8:
            raise RuntimeError("Le mot de passe doit contenir au moins 8 caractères.")
        existing["passwordHash"] = hash_password(password)

    data.setdefault("auditLog", []).insert(0, {
        "date": today_iso(),
        "actor": actor,
        "action": f"Compte utilisateur '{username}' enregistré avec le rôle {role}.",
    })
    return existing


def role_default_username(role: str) -> str:
    return {
        "Admin RH": "admin",
        "Assistant RH": "assistant",
        "Direction": "direction",
    }.get(normalize_role(role), "admin")


def find_login_user(data: dict, username: str, role: str) -> dict | None:
    username_key = normalize_lookup(username)
    role = normalize_role(role)
    for user in data.get("users", []):
        if user.get("active", True) is False:
            continue
        user_role = normalize_role(user.get("role"))
        user_key = normalize_lookup(user.get("username"))
        role_alias = normalize_lookup(role_default_username(user_role))
        role_key = normalize_lookup(user_role)
        if username_key and username_key in {user_key, role_alias, role_key} and (not role or user_role == role):
            return user
    return None


def authenticate_admin_login(username: str, role: str, password: str) -> tuple[str | None, dict | None]:
    role = normalize_role(role)
    data = get_state()
    username = str(username or "").strip()
    user = find_login_user(data, username or role_default_username(role), role)
    if user and verify_password_hash(password, user.get("passwordHash", "")):
        user["lastLoginAt"] = datetime.now().isoformat(timespec="seconds")
        save_state(data)
        return normalize_role(user.get("role")), public_user(user)
    if user:
        return None, None

    if authenticate_role(role, password):
        return role, {
            "id": f"env-{normalize_lookup(role)}",
            "username": role_default_username(role),
            "fullName": role,
            "role": role,
            "active": True,
            "source": "Render",
        }
    return None, None


def configured_role_passwords() -> dict[str, str]:
    return {role: password for role, password in ROLE_PASSWORDS.items() if password}


def stored_user_auth_enabled() -> bool:
    try:
        return any(
            user.get("active", True) is not False and user.get("passwordHash")
            for user in get_state().get("users", [])
        )
    except Exception:
        return False


def auth_required() -> bool:
    return bool(configured_role_passwords()) or stored_user_auth_enabled()


def normalize_role(role: str | None) -> str:
    role = str(role or "Admin RH").strip()
    return role if role in ROLE_PERMISSIONS else "Admin RH"


def authenticate_role(role: str, password: str) -> bool:
    passwords = configured_role_passwords()
    if not passwords:
        return not stored_user_auth_enabled()
    role = normalize_role(role)
    expected = passwords.get(role)
    if not expected:
        return False
    return hmac.compare_digest(password, expected)


def role_has_permission(role: str | None, permission: str) -> bool:
    role = normalize_role(role)
    return permission in ROLE_PERMISSIONS.get(role, set())


def sign_session(payload: str) -> str:
    return hmac.new(APP_SESSION_SECRET.encode("utf-8"), payload.encode("utf-8"), hashlib.sha256).hexdigest()


def create_session_cookie(role: str = "Admin RH") -> str:
    role = normalize_role(role)
    issued_at = str(int(time.time()))
    payload = f"admin:{role}:{issued_at}"
    signature = sign_session(payload)
    token = base64.urlsafe_b64encode(f"{role}:{issued_at}:{signature}".encode("utf-8")).decode("ascii")
    flags = "HttpOnly; Path=/; SameSite=Lax; Max-Age=" + str(SESSION_MAX_AGE_SECONDS)
    if os.environ.get("RENDER"):
        flags += "; Secure"
    return f"{SESSION_COOKIE}={token}; {flags}"


def cookie_flags(max_age: int) -> str:
    flags = f"HttpOnly; Path=/; SameSite=Lax; Max-Age={max_age}"
    if os.environ.get("RENDER"):
        flags += "; Secure"
    return flags


def create_employee_session_cookie(employee_id: str) -> str:
    issued_at = str(int(time.time()))
    payload = f"employee:{employee_id}:{issued_at}"
    signature = sign_session(payload)
    token = base64.urlsafe_b64encode(f"{employee_id}:{issued_at}:{signature}".encode("utf-8")).decode("ascii")
    return f"{EMPLOYEE_SESSION_COOKIE}={token}; {cookie_flags(SESSION_MAX_AGE_SECONDS)}"


def parse_cookie(header: str | None) -> dict[str, str]:
    if not header:
        return {}
    cookies: dict[str, str] = {}
    for part in header.split(";"):
        if "=" not in part:
            continue
        key, value = part.strip().split("=", 1)
        cookies[key] = value
    return cookies


def admin_role_from_cookie(header: str | None) -> str | None:
    if not auth_required():
        return "Admin RH"
    token = parse_cookie(header).get(SESSION_COOKIE)
    if not token:
        return None
    try:
        decoded = base64.urlsafe_b64decode(token.encode("ascii")).decode("utf-8")
        parts = decoded.split(":")
        if len(parts) == 2:
            issued_at, signature = parts
            if not hmac.compare_digest(signature, sign_session(issued_at)):
                return None
            role = "Admin RH"
        elif len(parts) == 3:
            role, issued_at, signature = parts
            role = normalize_role(role)
            payload = f"admin:{role}:{issued_at}"
            if not hmac.compare_digest(signature, sign_session(payload)):
                return None
        else:
            return None
        if time.time() - int(issued_at) > SESSION_MAX_AGE_SECONDS:
            return None
        return role
    except Exception:
        return None


def verify_session_cookie(header: str | None) -> bool:
    return bool(admin_role_from_cookie(header))


def verify_employee_session_cookie(header: str | None) -> str | None:
    token = parse_cookie(header).get(EMPLOYEE_SESSION_COOKIE)
    if not token:
        return None
    try:
        decoded = base64.urlsafe_b64decode(token.encode("ascii")).decode("utf-8")
        employee_id, issued_at, signature = decoded.split(":", 2)
        payload = f"employee:{employee_id}:{issued_at}"
        if not hmac.compare_digest(signature, sign_session(payload)):
            return None
        if time.time() - int(issued_at) > SESSION_MAX_AGE_SECONDS:
            return None
        return employee_id
    except Exception:
        return None


def rows(conn: sqlite3.Connection, sql: str, params: tuple = ()) -> list[sqlite3.Row]:
    return list(conn.execute(sql, params).fetchall())


def one(conn: sqlite3.Connection, sql: str, params: tuple = ()) -> sqlite3.Row | None:
    return conn.execute(sql, params).fetchone()


def json_value(value, fallback):
    if value is None:
        return fallback
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return fallback


def get_postgres_state() -> dict:
    with connect_postgres() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT value FROM app_state WHERE key = %s", ("main",))
            row = cur.fetchone()
            if not row:
                return default_state()
            return normalize_state(row[0])


def save_postgres_state(data: dict) -> None:
    state = normalize_state(data)
    with connect_postgres() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO app_state(key, value, updated_at)
                VALUES (%s, %s::jsonb, NOW())
                ON CONFLICT (key)
                DO UPDATE SET value = EXCLUDED.value, updated_at = NOW()
                """,
                ("main", json.dumps(state, ensure_ascii=False)),
            )


def create_database_backup() -> tuple[str, str]:
    backup_dir = backup_root()
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if DB_MODE == "postgres":
        backup_path = backup_dir / f"rh_control_backup_{stamp}.json"
        backup_path.write_text(json.dumps(get_state(), ensure_ascii=False, indent=2), encoding="utf-8")
        return str(backup_path), "json"

    backup_path = backup_dir / f"rh_control_backup_{stamp}.sqlite"
    if DB_PATH.exists():
        shutil.copy2(DB_PATH, backup_path)
    return str(backup_path), "sqlite"


def get_state() -> dict:
    if DB_MODE == "postgres":
        return get_postgres_state()

    with connect() as conn:
        settings_rows = rows(conn, "SELECT key, value FROM settings")
        settings = {item["key"]: json_value(item["value"], item["value"]) for item in settings_rows}

        employees = []
        for emp in rows(conn, "SELECT * FROM employees ORDER BY sort_order, last_name, first_name"):
            balance = one(conn, "SELECT * FROM leave_balances WHERE employee_id = ?", (emp["id"],))
            contracts = []
            for contract in rows(conn, "SELECT * FROM contracts WHERE employee_id = ? ORDER BY sort_order", (emp["id"],)):
                history = [
                    item["entry"]
                    for item in rows(conn, "SELECT entry FROM contract_history WHERE contract_id = ? ORDER BY sort_order, id", (contract["id"],))
                ]
                contracts.append(
                    {
                        "id": contract["id"],
                        "type": contract["type"],
                        "start": contract["start_date"] or "",
                        "end": contract["end_date"] or "",
                        "duration": contract["duration"] or "",
                        "fonction": contract["fonction"] or "",
                        "service": contract["service"] or "",
                        "salary": contract["salary"] or 0,
                        "renewalDate": contract["renewal_date"] or "",
                        "renewalCount": contract["renewal_count"] or 0,
                        "document": contract["document"] or "",
                        "status": contract["status"] or "",
                        "history": history,
                    }
                )

            employees.append(
                {
                    "id": emp["id"],
                    "matricule": emp["matricule"],
                    "firstName": emp["first_name"],
                    "lastName": emp["last_name"],
                    "service": emp["service"],
                    "direction": emp["direction_name"] or "",
                    "agency": emp["agency"] or "",
                    "fonction": emp["fonction"] or "",
                    "status": emp["status"] or "Actif",
                    "leaveBalance": {
                        "initial": balance["initial_days"] if balance else 0,
                        "acquired": balance["acquired_days"] if balance else 0,
                        "taken": balance["taken_days"] if balance else 0,
                        "planned": balance["planned_days"] if balance else 0,
                        "available": balance["available_days"] if balance else 0,
                    },
                    "contracts": contracts,
                }
            )

        leave_requests = []
        for leave in rows(conn, "SELECT * FROM leave_requests ORDER BY sort_order, created_at DESC"):
            observations = [
                item["entry"]
                for item in rows(conn, "SELECT entry FROM leave_observations WHERE leave_id = ? ORDER BY sort_order, id", (leave["id"],))
            ]
            history = [
                item["entry"]
                for item in rows(conn, "SELECT entry FROM leave_history WHERE leave_id = ? ORDER BY sort_order, id", (leave["id"],))
            ]
            leave_requests.append(
                {
                    "id": leave["id"],
                    "employeeId": leave["employee_id"],
                    "type": leave["type"],
                    "start": leave["start_date"] or "",
                    "end": leave["end_date"] or "",
                    "returnDate": leave["return_date"] or "",
                    "days": leave["days"] or 0,
                    "reason": leave["reason"] or "",
                    "comment": leave["comment"] or "",
                    "attachment": leave["attachment"] or "",
                    "status": leave["status"] or "",
                    "observations": observations,
                    "createdAt": leave["created_at"] or "",
                    "history": history,
                }
            )

        documents = [
            {
                "id": doc["id"],
                "employeeId": doc["employee_id"] or "",
                "leaveId": doc["leave_id"] or "",
                "title": doc["title"],
                "status": doc["status"] or "",
                "createdAt": doc["created_at"] or "",
                "content": doc["content"] or "",
            }
            for doc in rows(conn, "SELECT * FROM documents ORDER BY sort_order, created_at DESC")
        ]

        audit_log = [
            {
                "date": item["date"],
                "actor": item["actor"],
                "action": item["action"],
            }
            for item in rows(conn, "SELECT * FROM audit_log ORDER BY sort_order, id")
        ]

        default_settings = dict(DEFAULT_SETTINGS)
        default_settings.update(settings)

        role_row = one(conn, "SELECT value FROM app_meta WHERE key = ?", ("current_role",))
        full_state_row = one(conn, "SELECT value FROM app_meta WHERE key = ?", ("full_state_json",))
        if full_state_row:
            return normalize_state(full_state_row["value"])

        return {
            "currentRole": role_row["value"] if role_row else "Admin RH",
            "settings": default_settings,
            "employees": employees,
            "leaveRequests": leave_requests,
            "documentRequests": [],
            "documents": documents,
            "documentTemplates": dict(DEFAULT_DOCUMENT_TEMPLATES),
            "notifications": [],
            "personnelActions": [],
            "users": [],
            "auditLog": audit_log,
        }


def save_state(data: dict) -> None:
    if DB_MODE == "postgres":
        save_postgres_state(data)
        return

    with connect() as conn:
        conn.execute("BEGIN")
        for table in [
            "contract_history",
            "leave_observations",
            "leave_history",
            "documents",
            "leave_requests",
            "contracts",
            "leave_balances",
            "employees",
            "settings",
            "audit_log",
        ]:
            conn.execute(f"DELETE FROM {table}")

        conn.execute(
            "INSERT OR REPLACE INTO app_meta(key, value) VALUES (?, ?)",
            ("current_role", data.get("currentRole", "Admin RH")),
        )

        for key, value in data.get("settings", {}).items():
            conn.execute(
                "INSERT INTO settings(key, value) VALUES (?, ?)",
                (key, json.dumps(value, ensure_ascii=False)),
            )

        for emp_index, employee in enumerate(data.get("employees", [])):
            conn.execute(
                """
                INSERT INTO employees(
                  id, matricule, first_name, last_name, service, direction_name,
                  agency, fonction, status, sort_order
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    employee.get("id"),
                    employee.get("matricule", ""),
                    employee.get("firstName", ""),
                    employee.get("lastName", ""),
                    employee.get("service", ""),
                    employee.get("direction", ""),
                    employee.get("agency", ""),
                    employee.get("fonction", ""),
                    employee.get("status", "Actif"),
                    emp_index,
                ),
            )

            balance = employee.get("leaveBalance", {})
            conn.execute(
                """
                INSERT INTO leave_balances(
                  employee_id, initial_days, acquired_days, taken_days, planned_days, available_days
                ) VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    employee.get("id"),
                    balance.get("initial", 0),
                    balance.get("acquired", 0),
                    balance.get("taken", 0),
                    balance.get("planned", 0),
                    balance.get("available", 0),
                ),
            )

            for contract_index, contract in enumerate(employee.get("contracts", [])):
                conn.execute(
                    """
                    INSERT INTO contracts(
                      id, employee_id, type, start_date, end_date, duration, fonction,
                      service, salary, renewal_date, renewal_count, document, status, sort_order
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        contract.get("id"),
                        employee.get("id"),
                        contract.get("type", ""),
                        contract.get("start", ""),
                        contract.get("end", ""),
                        contract.get("duration", ""),
                        contract.get("fonction", ""),
                        contract.get("service", ""),
                        contract.get("salary", 0),
                        contract.get("renewalDate", ""),
                        contract.get("renewalCount", 0),
                        contract.get("document", ""),
                        contract.get("status", ""),
                        contract_index,
                    ),
                )
                for hist_index, entry in enumerate(contract.get("history", [])):
                    conn.execute(
                        "INSERT INTO contract_history(contract_id, entry, sort_order) VALUES (?, ?, ?)",
                        (contract.get("id"), entry, hist_index),
                    )

        for leave_index, leave in enumerate(data.get("leaveRequests", [])):
            conn.execute(
                """
                INSERT INTO leave_requests(
                  id, employee_id, type, start_date, end_date, return_date,
                  days, reason, comment, attachment, status, created_at, sort_order
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    leave.get("id"),
                    leave.get("employeeId"),
                    leave.get("type", ""),
                    leave.get("start", ""),
                    leave.get("end", ""),
                    leave.get("returnDate", ""),
                    leave.get("days", 0),
                    leave.get("reason", ""),
                    leave.get("comment", ""),
                    leave.get("attachment", ""),
                    leave.get("status", ""),
                    leave.get("createdAt", ""),
                    leave_index,
                ),
            )
            for obs_index, entry in enumerate(leave.get("observations", [])):
                conn.execute(
                    "INSERT INTO leave_observations(leave_id, entry, sort_order) VALUES (?, ?, ?)",
                    (leave.get("id"), entry, obs_index),
                )
            for hist_index, entry in enumerate(leave.get("history", [])):
                conn.execute(
                    "INSERT INTO leave_history(leave_id, entry, sort_order) VALUES (?, ?, ?)",
                    (leave.get("id"), entry, hist_index),
                )

        for doc_index, doc in enumerate(data.get("documents", [])):
            conn.execute(
                """
                INSERT INTO documents(
                  id, employee_id, leave_id, title, status, created_at, content, sort_order
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    doc.get("id"),
                    doc.get("employeeId") or None,
                    doc.get("leaveId") or None,
                    doc.get("title", ""),
                    doc.get("status", ""),
                    doc.get("createdAt", ""),
                    doc.get("content", ""),
                    doc_index,
                ),
            )

        for audit_index, item in enumerate(data.get("auditLog", [])):
            conn.execute(
                "INSERT INTO audit_log(date, actor, action, sort_order) VALUES (?, ?, ?, ?)",
                (
                    item.get("date", ""),
                    item.get("actor", ""),
                    item.get("action", ""),
                    audit_index,
                ),
            )

        conn.execute(
            "INSERT OR REPLACE INTO app_meta(key, value) VALUES (?, ?)",
            ("full_state_json", json.dumps(normalize_state(data), ensure_ascii=False)),
        )
        conn.commit()


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def send_json(self, payload: dict, status: int = 200, extra_headers: dict[str, str] | None = None) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        for key, value in (extra_headers or {}).items():
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(body)

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8")
        return json.loads(raw or "{}")

    def is_authenticated(self) -> bool:
        return verify_session_cookie(self.headers.get("Cookie"))

    def require_auth(self) -> bool:
        if self.is_authenticated():
            return True
        self.send_json({"ok": False, "error": "Authentification requise"}, status=401)
        return False

    def send_json_download(self, filename: str, payload: dict) -> None:
        body = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_binary_download(self, filename: str, body: bytes, content_type: str) -> None:
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def employee_id_from_cookie(self) -> str | None:
        return verify_employee_session_cookie(self.headers.get("Cookie"))

    def current_admin_role(self) -> str | None:
        return admin_role_from_cookie(self.headers.get("Cookie"))

    def require_permission(self, permission: str) -> str | None:
        role = self.current_admin_role()
        if not role:
            self.send_json({"ok": False, "error": "Authentification requise"}, status=401)
            return None
        if not role_has_permission(role, permission):
            self.send_json({"ok": False, "error": "Action non autorisée pour ce rôle"}, status=403)
            return None
        return role

    def require_employee(self) -> tuple[dict, dict] | None:
        employee_id = self.employee_id_from_cookie()
        if not employee_id:
            self.send_json({"ok": False, "error": "Connexion salarié requise"}, status=401)
            return None
        data = get_state()
        employee = find_employee_by_id(data, employee_id)
        if not employee:
            self.send_json({"ok": False, "error": "Matricule introuvable"}, status=401)
            return None
        return data, employee

    def do_GET(self):
        path = urlparse(self.path).path
        if path == "/api/health":
            self.send_json({
                "ok": True,
                "database": database_label(),
                "mode": DB_MODE,
                "durable": db_is_durable(),
                "authRequired": auth_required(),
                "emailConfigured": bool(SMTP_HOST and SMTP_FROM),
            })
            return
        if path == "/api/auth/status":
            role = self.current_admin_role()
            self.send_json({
                "ok": True,
                "authRequired": auth_required(),
                "authenticated": bool(role),
                "role": role,
            })
            return
        if path == "/api/employee/status":
            employee_id = self.employee_id_from_cookie()
            data = get_state()
            employee = find_employee_by_id(data, employee_id) if employee_id else None
            self.send_json({
                "ok": True,
                "authenticated": bool(employee),
                "employee": employee,
            })
            return
        if path == "/api/bootstrap":
            if not self.require_auth():
                return
            self.send_json({
                "ok": True,
                "database": database_label(),
                "durable": db_is_durable(),
                "role": self.current_admin_role(),
                "emailConfigured": bool(SMTP_HOST and SMTP_FROM),
                "data": state_for_client(get_state()),
            })
            return
        if path == "/api/users":
            if not self.require_permission("users_manage"):
                return
            data = get_state()
            self.send_json({
                "ok": True,
                "users": [public_user(user) for user in data.get("users", [])],
                "fallbackRoles": list(configured_role_passwords().keys()),
            })
            return
        if path == "/api/export":
            if not self.require_permission("export_json"):
                return
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.send_json_download(f"rh_control_export_{stamp}.json", get_state())
            return
        if path == "/api/export-excel":
            if not self.require_permission("export_excel"):
                return
            query = urlparse(self.path).query
            export_type = "summary"
            for part in query.split("&"):
                if part.startswith("type="):
                    export_type = part.split("=", 1)[1]
            try:
                body = build_export_workbook(export_type)
                stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                self.send_binary_download(
                    f"rapport_rh_{export_type}_{stamp}.xlsx",
                    body,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/employees/template":
            if not self.require_permission("excel_template"):
                return
            try:
                body = build_employee_template_workbook()
                self.send_binary_download(
                    "modele_import_collaborateurs.xlsx",
                    body,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/employee/bootstrap":
            result = self.require_employee()
            if not result:
                return
            data, employee = result
            self.send_json({"ok": True, "data": employee_portal_payload(data, employee)})
            return
        return super().do_GET()

    def do_POST(self):
        path = urlparse(self.path).path
        if path == "/api/auth/login":
            payload = self.read_json()
            password = payload.get("password", "")
            role = normalize_role(payload.get("role", "Admin RH"))
            username = str(payload.get("username", "")).strip()
            authenticated_role, user = authenticate_admin_login(username, role, password)
            if authenticated_role:
                self.send_json(
                    {"ok": True, "authRequired": auth_required(), "role": authenticated_role, "user": user},
                    extra_headers={"Set-Cookie": create_session_cookie(authenticated_role)},
                )
                return
            self.send_json({"ok": False, "error": "Mot de passe incorrect"}, status=401)
            return
        if path == "/api/auth/logout":
            self.send_json(
                {"ok": True},
                extra_headers={"Set-Cookie": f"{SESSION_COOKIE}=; HttpOnly; Path=/; SameSite=Lax; Max-Age=0"},
            )
            return
        if path == "/api/employee/login":
            payload = self.read_json()
            matricule = str(payload.get("matricule", "")).strip()
            data = get_state()
            employee = find_employee_by_matricule(data, matricule)
            if not employee:
                self.send_json({"ok": False, "error": "Matricule introuvable"}, status=401)
                return
            self.send_json(
                {"ok": True, "data": employee_portal_payload(data, employee)},
                extra_headers={"Set-Cookie": create_employee_session_cookie(employee["id"])},
            )
            return
        if path == "/api/employee/logout":
            self.send_json(
                {"ok": True},
                extra_headers={"Set-Cookie": f"{EMPLOYEE_SESSION_COOKIE}=; HttpOnly; Path=/; SameSite=Lax; Max-Age=0"},
            )
            return
        if path == "/api/employee/leave":
            result = self.require_employee()
            if not result:
                return
            data, employee = result
            payload = self.read_json()
            start = str(payload.get("start", "")).strip()
            end = str(payload.get("end", "")).strip()
            days = calc_leave_days_server(data.get("settings", {}), start, end)
            if not start or not end or days <= 0:
                self.send_json({"ok": False, "error": "Dates de congé invalides"}, status=400)
                return
            available = float(employee.get("leaveBalance", {}).get("available", 0) or 0)
            if days > available and not data.get("settings", {}).get("allowExceptionalLeave", True):
                self.send_json({"ok": False, "error": "Solde de congé insuffisant"}, status=400)
                return
            leave = {
                "id": f"leave-{uuid.uuid4().hex[:12]}",
                "employeeId": employee["id"],
                "type": str(payload.get("type", "Congé annuel")).strip() or "Congé annuel",
                "start": start,
                "end": end,
                "returnDate": calc_return_date_server(data.get("settings", {}), end),
                "days": days,
                "reason": str(payload.get("reason", "")).strip(),
                "comment": str(payload.get("comment", "")).strip(),
                "attachment": str(payload.get("attachment", "")).strip(),
                "status": "Demande envoyée",
                "observations": [],
                "createdAt": datetime.now().date().isoformat(),
                "history": [f"{datetime.now():%d/%m/%Y} : demande envoyée par le collaborateur."],
            }
            data.setdefault("leaveRequests", []).insert(0, leave)
            data.setdefault("auditLog", []).insert(0, {
                "date": datetime.now().date().isoformat(),
                "actor": "Collaborateur",
                "action": f"{employee.get('firstName', '')} {employee.get('lastName', '')} a envoyé une demande de congé.",
            })
            record_notification(
                data,
                "Congés",
                "Nouvelle demande de congé",
                f"{full_name_server(employee)} a demandé {days} jour(s), du {format_date_fr(start)} au {format_date_fr(end)}.",
                notification_recipients(data, "rh"),
            )
            save_state(data)
            refreshed = get_state()
            refreshed_employee = find_employee_by_id(refreshed, employee["id"]) or employee
            self.send_json({"ok": True, "data": employee_portal_payload(refreshed, refreshed_employee)})
            return
        if path == "/api/employee/document-request":
            result = self.require_employee()
            if not result:
                return
            data, employee = result
            payload = self.read_json()
            request_type = str(payload.get("type", "")).strip() or "Autres"
            details = str(payload.get("details", "")).strip()
            request = {
                "id": f"docreq-{uuid.uuid4().hex[:12]}",
                "employeeId": employee["id"],
                "type": request_type,
                "details": details,
                "status": "Demande envoyée",
                "createdAt": datetime.now().date().isoformat(),
                "content": "",
                "history": [f"{datetime.now():%d/%m/%Y} : demande envoyée par le salarié."],
            }
            data.setdefault("documentRequests", []).insert(0, request)
            data.setdefault("auditLog", []).insert(0, {
                "date": datetime.now().date().isoformat(),
                "actor": "Collaborateur",
                "action": f"Demande de document '{request_type}' envoyée par {employee.get('matricule', '')}.",
            })
            record_notification(
                data,
                "Documents",
                "Nouvelle demande de document RH",
                f"{full_name_server(employee)} a demandé : {request_type}. Précision : {details or '—'}.",
                notification_recipients(data, "rh"),
            )
            save_state(data)
            refreshed = get_state()
            refreshed_employee = find_employee_by_id(refreshed, employee["id"]) or employee
            self.send_json({"ok": True, "data": employee_portal_payload(refreshed, refreshed_employee)})
            return
        if path == "/api/admin/leave-transition":
            role = self.current_admin_role()
            if not role:
                self.send_json({"ok": False, "error": "Authentification requise"}, status=401)
                return
            try:
                payload = self.read_json()
                data = get_state()
                transition_leave_server(
                    data,
                    str(payload.get("id", "")).strip(),
                    str(payload.get("action", "")).strip(),
                    role,
                    str(payload.get("observation", "")).strip(),
                )
                save_state(data)
                self.send_json({"ok": True, "data": state_for_client(get_state())})
            except PermissionError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=403)
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if path == "/api/documents/render":
            if not self.require_permission("document_render"):
                return
            try:
                payload = self.read_json()
                body, filename, content_type = render_document_file(
                    str(payload.get("title", "Document RH")),
                    str(payload.get("content", "")),
                    str(payload.get("format", "docx")),
                )
                self.send_binary_download(filename, body, content_type)
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/templates/import-docx":
            if not self.require_permission("template_import"):
                return
            try:
                length = int(self.headers.get("Content-Length", "0"))
                body = self.rfile.read(length)
                content = extract_multipart_file(body, self.headers.get("Content-Type", ""))
                text = import_docx_template(content)
                self.send_json({"ok": True, "content": text})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/users":
            role = self.require_permission("users_manage")
            if not role:
                return
            try:
                payload = self.read_json()
                data = get_state()
                upsert_user(data, payload, role)
                save_state(data)
                refreshed = get_state()
                self.send_json({"ok": True, "users": [public_user(user) for user in refreshed.get("users", [])]})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if path == "/api/employees/preview-excel":
            if not self.require_permission("excel_import"):
                return
            try:
                length = int(self.headers.get("Content-Length", "0"))
                body = self.rfile.read(length)
                content = extract_multipart_file(body, self.headers.get("Content-Type", ""))
                result = preview_employees_from_excel(content)
                self.send_json({"ok": True, **result})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/employees/import-excel":
            role = self.require_permission("excel_import")
            if not role:
                return
            try:
                length = int(self.headers.get("Content-Length", "0"))
                body = self.rfile.read(length)
                content = extract_multipart_file(body, self.headers.get("Content-Type", ""))
                result = import_employees_from_excel(content, actor=role)
                result["data"] = state_for_client(result.get("data", get_state()))
                self.send_json({"ok": True, **result})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/backup":
            if not self.require_permission("backup"):
                return
            try:
                backup_path, backup_type = create_database_backup()
                self.send_json({"ok": True, "backup": backup_path, "type": backup_type, "durable": db_is_durable()})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        self.send_json({"ok": False, "error": "Route inconnue"}, status=404)

    def do_PUT(self):
        path = urlparse(self.path).path
        if path == "/api/state":
            role = self.require_permission("state_write")
            if not role:
                return
            try:
                payload = self.read_json()
                save_state(sanitize_state_for_role(payload.get("data", payload), role))
                self.send_json({"ok": True, "database": database_label()})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        if path == "/api/import":
            if not self.require_permission("import_json"):
                return
            try:
                payload = self.read_json()
                data = payload.get("data", payload)
                save_state(data)
                self.send_json({"ok": True, "database": database_label(), "data": state_for_client(get_state())})
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=500)
            return
        self.send_json({"ok": False, "error": "Route inconnue"}, status=404)


def main() -> None:
    init_db()
    display_host = "127.0.0.1" if HOST == "0.0.0.0" else HOST
    url = f"http://{display_host}:{PORT}/"
    print("Application RH démarrée")
    print(f"URL        : {url}")
    print(f"Base       : {database_label()}")
    if not os.environ.get("PORT"):
        print("Laisse cette fenêtre ouverte pendant l'utilisation.")
    if "--no-browser" not in sys.argv and not os.environ.get("PORT"):
        try:
            webbrowser.open(url)
        except Exception:
            pass
    ThreadingHTTPServer((HOST, PORT), Handler).serve_forever()


if __name__ == "__main__":
    main()
